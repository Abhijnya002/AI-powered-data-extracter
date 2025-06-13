# from http.server import BaseHTTPRequestHandler
# import json
# import os
# import tempfile
# from urllib.parse import parse_qs
# from io import BytesIO

# import pandas as pd
# import docx
# import re
# from openpyxl import Workbook
# from openpyxl.styles import PatternFill, Font
# from transformers import pipeline
# import torch
# from base64 import b64decode

# summarizer = pipeline("summarization", model="sshleifer/distilbart-cnn-12-6", device=0 if torch.cuda.is_available() else -1)
# title_gen = pipeline("text2text-generation", model="google/flan-t5-base", device=0 if torch.cuda.is_available() else -1)

# class handler(BaseHTTPRequestHandler):
#     def do_POST(self):
#         length = int(self.headers.get('content-length'))
#         data = self.rfile.read(length)

#         boundary = self.headers.get("content-type").split("boundary=")[1]
#         parts = data.split(boundary.encode())[1:-1]
#         content = parts[0].split(b"\r\n\r\n")[1].rstrip(b"\r\n--")

#         with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as temp_file:
#             temp_file.write(content)
#             temp_file_path = temp_file.name

#         # Process DOCX
#         try:
#             doc = docx.Document(temp_file_path)
#             text = re.sub(r"\s+", " ", " ".join(p.text for p in doc.paragraphs)).strip()

#             # Extract task sentences (simple filter)
#             tasks = [s for s in re.split(r'(?<=[.!?])\s+', text) if len(s.split()) > 4]

#             output = {
#                 "summary": summarizer(tasks[0], max_length=60, min_length=20, do_sample=False)[0]["summary_text"],
#                 "title": title_gen(f"Create a concise renovation title from: \"{tasks[0]}\"", max_length=20, do_sample=False)[0]["generated_text"]
#             }

#             self.send_response(200)
#             self.send_header("Content-Type", "application/json")
#             self.end_headers()
#             self.wfile.write(json.dumps(output).encode())

#         except Exception as e:
#             self.send_response(500)
#             self.end_headers()
#             self.wfile.write(str(e).encode())
from transformers import pipeline
import torch
import docx
import tempfile
import json
import re

def handler(request):
    try:
        content_type = request.headers.get("content-type", "")
        if "multipart/form-data" not in content_type:
            return {
                "statusCode": 400,
                "body": json.dumps({"error": "Invalid content type"})
            }

        boundary = content_type.split("boundary=")[-1]
        body = request.body
        parts = body.split(boundary.encode())[1:-1]
        content = parts[0].split(b"\r\n\r\n")[1].rstrip(b"\r\n--")

        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as temp_file:
            temp_file.write(content)
            temp_file_path = temp_file.name

        # Extract text from docx
        doc = docx.Document(temp_file_path)
        text = re.sub(r"\s+", " ", " ".join(p.text for p in doc.paragraphs)).strip()
        tasks = [s for s in re.split(r'(?<=[.!?])\s+', text) if len(s.split()) > 4]

        if not tasks:
            return {
                "statusCode": 400,
                "body": json.dumps({"error": "No valid task sentences found."})
            }

        # ✅ Move model loading INSIDE the function
        summarizer = pipeline("summarization", model="t5-small", device=-1)
        title_gen = pipeline("text2text-generation", model="t5-small", device=-1)

        summary = summarizer(tasks[0], max_length=60, min_length=20, do_sample=False)[0]["summary_text"]
        title = title_gen(f"Create a concise renovation title from: \"{tasks[0]}\"", max_length=20, do_sample=False)[0]["generated_text"]

        return {
            "statusCode": 200,
            "headers": {"Content-Type": "application/json"},
            "body": json.dumps({
                "summary": summary,
                "title": title
            })
        }

    except Exception as e:
        return {
            "statusCode": 500,
            "body": json.dumps({"error": str(e)})
        }
