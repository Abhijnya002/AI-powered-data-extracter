# import os
# import pandas as pd
# import docx
# import re
# from openpyxl import Workbook
# from openpyxl.styles import PatternFill, Font
# from transformers import pipeline
# import torch

# # --- Local Model Setup with Lighter Models ---
# summarizer = pipeline("summarization", model="sshleifer/distilbart-cnn-12-6", device=0 if torch.cuda.is_available() else -1)
# title_gen = pipeline("text2text-generation", model="google/flan-t5-small", device=0 if torch.cuda.is_available() else -1)

# # --- File Processing ---
# def load_transcript_text(docx_path):
#     try:
#         doc = docx.Document(docx_path)
#     except Exception as e:
#         raise ValueError("Failed to read the .docx file. Make sure it's valid.") from e
#     return re.sub(r"\s+", " ", " ".join(p.text for p in doc.paragraphs)).strip()

# def extract_task_sentences(text):
#     keywords = [
#         "washer", "dryer", "cabinet", "countertop", "floor", "door", "hvac", "vent", "appliance",
#         "bathroom", "kitchen", "electrical", "panel", "replace", "install", "refinish", "upgrade",
#         "renovate", "remodel", "materials", "tile", "riser", "drawing", "sketch", "plan"
#     ]
#     return [s for s in re.split(r'(?<=[.!?])\s+', text)
#             if any(k in s.lower() for k in keywords) and len(s.split()) > 3]

# def infer_category(sentence):
#     s = sentence.lower()
#     if "kitchen" in s: return "Kitchen"
#     if "bathroom" in s: return "Bathroom"
#     if "hvac" in s or "ac" in s: return "HVAC"
#     if "electrical" in s or "panel" in s: return "Electrical"
#     if "floor" in s: return "Flooring"
#     if "door" in s: return "Doors"
#     if "washer" in s or "dryer" in s: return "Laundry"
#     if "tile" in s: return "Tiling"
#     return "General"

# def extract_all_budgets(text):
#     numeric = re.findall(r"(?:\$|USD\s?)?(\d{1,3}(?:,\d{3})+|\d+)(?:\s?(?:dollars|bucks))?", text)
#     informal = re.findall(r"(\d+)\s?(?:grand|grands)", text, flags=re.I)
#     cleaned_numeric = []
#     for n in numeric:
#         try:
#             val = int(n.replace(",", ""))
#             if val > 100:
#                 cleaned_numeric.append(val)
#         except:
#             continue
#     return cleaned_numeric + [int(i) * 1000 for i in informal]

# def extract_lead(sentence):
#     s = sentence.lower()
#     if "al will" in s or "al is going to" in s: return "Al"
#     if "you will" in s or "send to you" in s: return "Contractor"
#     if "he will" in s or "she will" in s or "they will" in s: return "Client or Team"
#     if "designer" in s: return "Designer"
#     return ""

# def extract_drawing_ref(sentence):
#     s = sentence.lower()
#     if "architectural" in s and "drawing" in s:
#         return "Architectural drawing"
#     elif "riser" in s:
#         return "Plumbing riser detail"
#     elif "plan" in s:
#         return "Renovation plan"
#     elif "sketch" in s:
#         return "Design sketch"
#     elif "drawing" in s:
#         return "General drawing reference"
#     return ""

# def clean_title(text):
#     return re.sub(r"^(how|what|when|why|where|who)\b.*", "", text, flags=re.I).strip().rstrip(":?.")

# # --- AI Processing ---
# def process_tasks(sentences, budgets, file_title):
#     grouped = {}
#     used = []

#     for s in sentences:
#         if "?" in s or len(s.strip()) < 25: continue
#         category = infer_category(s)

#         raw_title = title_gen(
#             f"Create a clear and concise renovation task title from: \"{s}\"",
#             max_length=20, do_sample=False
#         )[0]["generated_text"]
#         title = clean_title(raw_title)

#         proposed = summarizer(s, max_length=60, min_length=15, do_sample=False)[0]["summary_text"]

#         comment_prompt = f"Explain this renovation task in more detail with extra clarity:\n\n\"{s}\""
#         comment = title_gen(comment_prompt, max_length=60, do_sample=False)[0]["generated_text"]

#         budget = 0
#         if budgets:
#             budget = budgets.pop(0)
#             used.append(budget)

#         task = {
#             "Task Description": title or "Renovation Task",
#             "Budget": budget,
#             "Proposed": proposed,
#             "Comment": comment,
#             "Drawing Ref": extract_drawing_ref(s),
#             "Lead": extract_lead(s)
#         }
#         grouped.setdefault(category, []).append(task)

#     return grouped, sum(used)

# # --- Save to Excel ---
# def save_structured_excel(grouped_tasks, total_budget, output_path, file_title):
#     file_base = os.path.basename(file_title)
#     wb = Workbook()
#     ws = wb.active
#     ws.title = "Scope of Work"
#     ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=7)
#     ws["A1"] = f"Scope of Work: {file_base}"
#     ws["A1"].font = Font(bold=True, size=14)
#     row = 3
#     cat_fill = PatternFill(start_color="D9EAD3", end_color="D9EAD3", fill_type="solid")

#     for category, tasks in grouped_tasks.items():
#         ws.cell(row=row, column=1, value=category)
#         for col in range(1, 8):
#             ws.cell(row=row, column=col).fill = cat_fill
#             ws.cell(row=row, column=col).font = Font(bold=True)
#         row += 1
#         for t in tasks:
#             ws.append(["", t["Task Description"], t["Budget"], t["Proposed"], t["Comment"],
#                        t["Drawing Ref"], t["Lead"]])
#             row += 1

#     ws.append(["TOTAL", "", total_budget, "", "", "", ""])
#     headers = ["Category", "Task Description", "Budget", "Proposed", "Comment", "Drawing Ref", "Lead"]
#     ws.insert_rows(3)
#     for i, h in enumerate(headers, 1):
#         ws.cell(row=3, column=i, value=h).font = Font(bold=True)

#     wb.save(output_path)
#     print(f"Excel saved: {output_path}")

# from flask import Flask, request, send_file
# from flask_cors import CORS
# import traceback

# app = Flask(__name__)
# CORS(app)

# @app.route('/api/process-docx', methods=['POST'])
# def process_docx():
#     try:
#         print("\U0001F4E9 Request received at /api/process-docx")

#         file = request.files['file']
#         print(f"\U0001F4E5 Received file: {file.filename}")
#         input_path = f"temp/{file.filename}"
#         os.makedirs(os.path.dirname(input_path), exist_ok=True)
#         file.save(input_path)
#         print(f"\U0001F4C1 File saved to: {input_path}")

#         transcript = load_transcript_text(input_path)
#         print(f"\U0001F4C4 Extracted transcript (first 100 chars): {transcript[:100]}")

#         sentences = extract_task_sentences(transcript)
#         print(f"\U0001F4DD Extracted {len(sentences)} task sentences")
#         budgets = extract_all_budgets(transcript)
#         print(f"\U0001F4B0 Extracted {len(budgets)} budgets")

#         grouped, total = process_tasks(sentences, budgets, input_path)
#         print(f"\U0001F4CA Processed tasks into {len(grouped)} categories. Total budget: ${total}")

#         output_path = input_path.replace('.docx', '.xlsx')
#         save_structured_excel(grouped, total, output_path, input_path)
#         print(f"✅ Excel file created at: {output_path}")

#         return send_file(output_path, as_attachment=True)

#     except Exception as e:
#         print("❌ Error occurred:", e)
#         traceback.print_exc()  # Add this line to print the stack trace
#         return f"Backend error: {str(e)}", 500

# if __name__ == '__main__':
#     port = int(os.environ.get("PORT", 10000))
#     app.run(host='0.0.0.0', port=port)
from flask import Flask, request, jsonify
import tempfile
import docx
import re
from transformers import pipeline
import torch

app = Flask(__name__)

@app.route('/process_doc', methods=['POST'])
def process_doc():
    try:
        file = request.files['file']
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as temp_file:
            file.save(temp_file.name)
            temp_file_path = temp_file.name

        doc = docx.Document(temp_file_path)
        text = re.sub(r"\s+", " ", " ".join(p.text for p in doc.paragraphs)).strip()
        tasks = [s for s in re.split(r'(?<=[.!?])\s+', text) if len(s.split()) > 4]

        if not tasks:
            return jsonify({"error": "No valid task sentences found."}), 400

        summarizer = pipeline("summarization", model="t5-small", device=-1)
        title_gen = pipeline("text2text-generation", model="t5-small", device=-1)

        summary = summarizer(tasks[0], max_length=60, min_length=20, do_sample=False)[0]["summary_text"]
        title = title_gen(f"Create a concise renovation title from: \"{tasks[0]}\"", max_length=20, do_sample=False)[0]["generated_text"]

        return jsonify({"summary": summary, "title": title})

    except Exception as e:
        return jsonify({"error": str(e)}), 500

app.run(host='0.0.0.0', port=8080)
