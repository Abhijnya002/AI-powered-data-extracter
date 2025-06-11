import os
import sys
from docx import Document
import pandas as pd
from openpyxl import Workbook

def process_docx(file_path):
    doc = Document(file_path)
    text = "\n".join([para.text for para in doc.paragraphs])
    
    # Your processing logic here
    # Example: Extract tasks and budgets
    tasks = [para.text for para in doc.paragraphs if len(para.text) > 20]
    
    # Create Excel output
    output_path = os.path.join(os.path.dirname(file_path), 'processed_output.xlsx')
    df = pd.DataFrame({'Tasks': tasks})
    df.to_excel(output_path, index=False)
    
    return output_path

if __name__ == '__main__':
    input_file = sys.argv[1]
    output_file = process_docx(input_file)
    print(output_file)  # Don't print anything else
